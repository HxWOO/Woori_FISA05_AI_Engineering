from pathlib import Path
import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.table import Table
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm
from step_1_1 import OUT_DIR
from step_1_2 import OUT_1_2
from step_3_1 import apply_font
from step_3_2 import add_blank_paragraph
from step_3_3 import OUT_3_3

OUT_3_4 = OUT_DIR / f"{Path(__file__).stem}.docx"

def insert_deposit_info(n_rows: int = 10):
    doc = Document(OUT_3_3)
    r_head = doc.add_paragraph().add_run("2. 주요 정기예금 상품 및 금리")
    apply_font(r_head, size_pt=14, is_bold=True)
    add_blank_paragraph(doc, size_pt=2)

    table = doc.add_table(rows=1, cols=6, style="Light Shading Accent 4")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    tr = table.rows[0]
    th_text = ["금융기관", "상품명", "이자계산", "만기(월)", "세전금리", "최고우대"]
    col_width = [Mm(40), Mm(53), Mm(20), Mm(20), Mm(20), Mm(20)]
    for idx, th in enumerate(tr.cells):
        th.width - col_width[idx]
        th.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p_th = th.paragraphs[0]
        p_th.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r_th = p_th.add_run(f"{th_text[idx]}")
        apply_font(r_th, size_pt=12, is_bold=True)

    df_raw = pd.read_excel(OUT_1_2)  # 정기예금 데이터
    df_filter = df_raw.filter(["kor_co_nm", "fin_prdt_nm", "intr_rate_type_nm", "save_trm", "intr_rate", "intr_rate2"])
    df_sort = df_filter.sort_values("intr_rate", ascending=False)

    for _, se_row in df_sort.head(n_rows).iterrows():
        tr = table.add_row()
        for idx, td in enumerate(tr.cells):
            td.width = col_width[idx]
            td.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p_td = td.paragraphs[0]
            if idx < 2:
                p_td.alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
            p_td.paragraph_format.space_before = Mm(2)  # 단락 앞 뒤 간격
            p_td.paragraph_format.space_after = Mm(2)
            p_td.add_run(f"{se_row.iloc[idx]}")

    add_blank_paragraph(doc, size_pt=10)
    doc.save(OUT_3_4)

if __name__ == "__main__":
    insert_deposit_info(20)