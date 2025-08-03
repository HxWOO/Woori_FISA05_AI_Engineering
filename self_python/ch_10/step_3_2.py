from datetime import datetime
from pathlib import Path
from docx import Document
from docx.document import Document as DocumentObject
from step_1_1 import OUT_DIR
from step_3_1 import OUT_3_1, apply_font

OUT_3_2 = OUT_DIR / f"{Path(__file__).stem}.docx"

def add_blank_paragraph(doc: DocumentObject, size_pt: int = None):
    r_empty = doc.add_paragraph().add_run(" ") # 빈 단락 삽입
    apply_font(r_empty, size_pt=size_pt)

def add_title(title:str = "정기예금 금리 현황표"):
    doc = Document(OUT_3_1)
    p_title = doc.add_paragraph(style="Title")
    r_title = p_title.add_run(title)
    apply_font(r_title, face="Malgun Gothic", size_pt=20, is_bold=True)
    now = datetime.now()  # 현재 시점 datetime 객체
    now_string = now.isoformat(sep=" ", timespec="minutes")  # 작성 일시
    r_now = p_title.add_run(f" (작성일시: {now_string})")
    apply_font(r_now, size_pt=14)
    add_blank_paragraph(doc, size_pt=6)
    doc.save(OUT_3_2)  # 워드 파일로 저장

if __name__ == "__main__":
    add_title()